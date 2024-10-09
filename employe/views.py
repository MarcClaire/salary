from django.shortcuts import render, redirect
from django.http import HttpResponseRedirect, HttpResponse
from django.urls import reverse
from .models import Employe, Contrat, Structure, Dossiers, TypeContrat,Paiement
from .forms import EmployeForm, ContratForm, StructureForum, DossiersFrom,TypeContratForm
from django.shortcuts import get_object_or_404
from elementSal.models import Elementsalaire
from docx import Document
#from docx2pdf import convert
import os, io
#from win32com import client as win32_client
#import pythoncom
import pdfkit
from django.conf import settings
from num2words import num2words
from datetime import timedelta
from calcul.Traitement import salaire_brut
from calcul.models import Parametre_calcul
# Create your views here.

from django.contrib.auth.models import User,Group
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Q
from SystemConf.Back_Control_Access import is_admin

# pip install fpdf
#from subprocess import run
#from fpdf import FPDF
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from django.utils import timezone
from datetime import timedelta
from django.db.models import Count
from django.db.models.functions import TruncYear


from django import forms




# @login_required(login_url="user_signin")
# def index(request):
#     total_employe=Employe.objects.all().count()
#     total_cdd=Contrat.objects.filter(type_contrat_id =1).all().count()
#     total_cdi=Contrat.objects.filter(type_contrat_id =2).count()
#     total_structure=Structure.objects.all().count()
#     return render(request, 'index.html', {'total_employe':total_employe, 'total_structure': total_structure,'total_cdd':total_cdd, 'total_cdi':total_cdi})

@login_required(login_url="user_signin")
def index(request):
    total_employe = Employe.objects.count()
    total_structure = Structure.objects.count()
    today = timezone.now().date()
    end_date_threshold = today + timedelta(days=30)
    # Interroger les contrats dont la date de fin est dans les 30 jours
    contrats = Contrat.objects.filter(date_fin__gte=today, date_fin__lte=end_date_threshold)
    context = {
        'total_employe': total_employe,
        'total_structure': total_structure,
        'contrats': contrats,
    }
    return render(request, 'dash.html', context)

#cette fonction permet d'affichier la liste des employé
@login_required(login_url="user_signin")
def liste(request):
    employe_object = Employe.objects.all()
    return render(request,'employe/list.html',{'employe_object':employe_object})

@login_required(login_url="user_signin")
def form(request):
    return render(request, 'employe/ajout_employe.html')

#Cette fonction permet d'ajouter un employé
@login_required(login_url="user_signin")
def ajouter_employe(request):
    if request.method == 'POST':
        form = EmployeForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                form.save()
                # nom_employe = employe.nom
                return HttpResponseRedirect(reverse('add_contrat'))
            except forms.ValidationError as e:
                # Si une ValidationError est levée, renvoyer le formulaire avec l'erreur
                return render(request, 'employe/ajout_employe.html', {'form': form, 'error_message': str(e)})
    else:
        form = EmployeForm()
    return render(request, 'employe/ajout_employe.html', {'form': form})

#Cette fonction permet de modier les informations d'un employé
@login_required(login_url="user_signin")
def update_empl_data(request, id):
    if request.method== 'POST':
        employe= Employe.objects.get(pk=id)
        form = EmployeForm(request.POST,request.FILES, instance=employe)
        if form.is_valid():
            try:
                form.save()
                

                return HttpResponseRedirect(reverse('personel'))
            except forms.ValidationError as e:
                # Si une ValidationError est levée, renvoyer le formulaire avec l'erreur
                return render(request, 'employe/ajout_employe.html', {'form': form, 'error_message': str(e)})
    else:
        employe = Employe.objects.get(pk=id)
        form = EmployeForm(instance=employe)
    return render(request, 'employe/update_employe.html', {'form':form} )



#Cette fonction permet de supprmer un employé
@login_required(login_url="user_signin")
def delete_employe(request, id):
    # if request.method=='POST':
    employe = Employe.objects.get(pk=id)
    employe.delete()
    return HttpResponseRedirect(reverse('personel'))

#cette fonction permetd'afficher la liste des contrat
@login_required(login_url="user_signin")
def liste_contrat(request):
    contrats = Contrat.objects.all()
    return render(request,'contrat/liste_contrat.html', {'contrats': contrats})

#Cette fonction permt d'ajouter un contrat
@login_required(login_url="user_signin")
def ajouter_contrat(request):
    if request.method == 'POST':
        form = ContratForm(request.POST)
        if form.is_valid():
            contrat=form.save(commit=False)
            contrat.calculer_date_retraite()
            contrat.save()
        return HttpResponseRedirect(reverse('contrat'))
        #return redirect(' contrat')  
    else:
        form = ContratForm()
    return render(request, 'contrat/ajout_contrat.html', {'form': form})

#Cette fonction permet de supprmer un contrat
@login_required(login_url="user_signin")
def delete_contrat(request, id):
    #if request.method=='POST':
    contrat = Contrat.objects.get(pk=id)
    contrat.delete()
    return HttpResponseRedirect(reverse('contrat'))

@login_required(login_url="user_signin")
def update_contrat_data(request, contrat_id):
    contrat = get_object_or_404(Contrat, id=contrat_id)
    if request.method== 'POST':
        #contrat= Contrat.objects.get(pk=contrat_id)
        form = ContratForm(request.POST, instance=contrat)
        if form.is_valid():
            #form.save()
            contrats = form.save(commit=False)
            contrats.calculer_date_retraite()
            contrats.save()
        return HttpResponseRedirect(reverse('contrat'))
        #return redirect('contrat')
    else:
        #contrat = Contrat.objects.get(pk=contrat_id)
        form = ContratForm(instance=contrat)
    return render(request, 'contrat/update_contrat.html', {'form':form} )

@login_required(login_url="user_signin")
def detail_contrat(request, id):
    detailcontrats=get_object_or_404(Contrat,pk=id)
    return render(request,'contrat/show_contrat.html', {'detailcontrats':detailcontrats})


# gestion des différentes structures
@login_required(login_url="user_signin")
def liste_structure(request):
    structures = Structure.objects.all()
    return render(request,'structure/list_structure.html',{'structures':structures})

@login_required(login_url="user_signin")
def ajouter_structure(request):
    if request.method == 'POST':
        form = StructureForum(request.POST)
        if form.is_valid():
            form.save()
        return HttpResponseRedirect(reverse('structure'))
    else:
        form = StructureForum()
    return render(request, 'structure/ajout_structure.html', {'form': form})

#Cette fonction permet de supprmer une stucture
@login_required(login_url="user_signin")
def delete_structure(request, id):
    structure = Structure.objects.get(pk=id)
    structure.delete()
    return HttpResponseRedirect(reverse('structure'))

@login_required(login_url="user_signin")
def update_structure(request, id):
    if request.method== 'POST':
        structure= Structure.objects.get(pk=id)
        form = StructureForum(request.POST, instance=structure)
        if form.is_valid():
            form.save()
        return HttpResponseRedirect(reverse('structure'))
    else:
        structure = Structure.objects.get(pk=id)
        form = StructureForum(instance=structure)
    return render(request, 'structure/update_structure.html', {'form':form} )

@login_required(login_url="user_signin")
def detail_structure(request, id):
    detailstructures=get_object_or_404(Contrat,pk=id)
    return render(request,'structure/detail_structure.html', {'detailstructures':detailstructures})

@login_required(login_url="user_signin")
def liste_employe_structure(request, structure_id):
    structure = get_object_or_404(Structure, id=structure_id)
    contrats = Contrat.objects.filter(structure=structure)
    employes = [contrat.employe for contrat in contrats]
    nombre_employes = len(employes)
    return render(request, 'structure/liste_employe_structure.html', {'structure': structure, 'employes': employes, 'nombre_employes': nombre_employes})

@login_required(login_url="user_signin")
def get_contrats(request, employe_id):
    employe = Employe.objects.get(id=employe_id)
    contrats = employe.contrat_set.all()
    return render(request, 'employe/liste_contrat_employe.html', {'contrats': contrats})


#cette fonction permet d'afficher la liste des fichiers charger par l'employé
@login_required(login_url="user_signin")
def dossier(request, id):
    try:
        employe = Employe.objects.get(id=id)
    except Employe.DoesNotExist:
        return HttpResponse("L'employé spécifié n'existe pas.", status=404)
    employee_files = Dossiers.objects.filter(employe=employe)
    return render(request, 'dossiers/liste_dossiers_employe.html', {'files': employee_files, 'employe':employe})

#Cette fonction permet d'ajouter un nouveau fichier
@login_required(login_url="user_signin")
def ajouter_dossiers(request, id):
    if request.method == 'POST':
        libelle = request.POST.get('libelle')
        file = request.FILES.get('file')
        print(f'Libelle: {libelle}, File: {file}')
        employe = get_object_or_404(Employe, id=id)
        if libelle and file:
            dossier = Dossiers(employe=employe, libelle=libelle, file=file)
            dossier.save()
        return redirect('dossier_employe', id=employe.id)
    else:
        return render(request, 'dossiers/liste_dossiers_employe.html', {'employe_id': employe.id})

#cette fonction permet de mettre à jour les dossiers
@login_required(login_url="user_signin")
def update_dossier(request, id):
    dossier = get_object_or_404(Dossiers,id=id)
    employe_id=dossier.employe.id
    if request.method == 'POST':
        libelle = request.POST.get('libelle')
        file = request.FILES.get('file')
        if libelle and file:
            dossier.libelle = libelle
            dossier.file = file
            print(f'Libelle: {libelle}, File: {file}')
            dossier.save()
        return redirect('dossier_employe', id=employe_id)
    return render(request, 'dossiers/liste_dossiers_employe.html', {'dossier': dossier})


#Cette fonction permet de supprmer un un dossier
@login_required(login_url="user_signin")
def deletedossier(request, id):
    dossier = get_object_or_404(Dossiers, pk=id)
    employe_id = dossier.employe.id  # Obtenez l'identifiant de l'employé avant de supprimer le dossier
    dossier.delete()
    return redirect('dossier_employe', id=employe_id)

@login_required(login_url="user_signin")
def contrat_doc_list(request):
    contrats = ContratDoc.objects.all()
    return render(request, 'contrat_doc_list.html', {'contrats': contrats})

#cette fonction charge un contrat
@login_required(login_url="user_signin")
def upload_contrat_doc(request):
    if request.method == 'POST':
       #type_contrat = request.POST.get('type_contrat')
        libelle = request.POST.get('libelle')
        file = request.FILES.get('file')
        if libelle and file:
            contrat_dac=TypeContrat(libelle=libelle, file=file)
            contrat_dac.save()

        return redirect('index')
    else:
        form = ContratDocForm()
    return render(request, 'contrat/upload_contrat_doc.html', {'form': form})

@login_required(login_url="user_signin")
def get_employes(request, employe_id):
    employe = get_object_or_404(Employe, id=employe_id)
    return render(request, 'employe/show_employe.html', {'employe': employe})

#la création de la fonctionnalité de gestion des départs à la retraite.
def apercu_departs_retraite(request):
    aujourd_hui = timezone.now().date()
    un_an = aujourd_hui + timedelta(days=365)
    trois_ans = aujourd_hui + timedelta(days=3*365)
    cinq_ans = aujourd_hui + timedelta(days=5*365)
    dix_ans = aujourd_hui + timedelta(days=10*365)

    departs_1an = Contrat.objects.filter(date_retraite_prevue__lte=un_an).order_by('date_retraite_prevue')
    departs_3ans = Contrat.objects.filter(date_retraite_prevue__lte=trois_ans).order_by('date_retraite_prevue')
    departs_5ans = Contrat.objects.filter(date_retraite_prevue__lte=cinq_ans).order_by('date_retraite_prevue')
    departs = Contrat.objects.filter(date_retraite_prevue__range=[aujourd_hui, dix_ans]).annotate(
        annee=TruncYear('date_retraite_prevue')).values('annee').annotate(
        count=Count('id')).order_by('annee')
    departs_par_departement = Contrat.objects.filter(date_retraite_prevue__range=[aujourd_hui, dix_ans]).values('structure').annotate(
        count=Count('id')).order_by('-count')

    departements = [depart['structure'] for depart in departs_par_departement]
    data_par_departement = [depart['count'] for depart in departs_par_departement]



    labels = [depart['annee'].year for depart in departs]
    data = [depart['count'] for depart in departs]

    context = {
        'labels': labels,
        'data': data,
        'departs_1an': departs_1an,
        'departs_3ans': departs_3ans,
        'departs_5ans': departs_5ans,
        'departs_1àans':departs,
        'departements': departements,
        'data_par_departement': data_par_departement,
    }
    return render(request, 'contrat/apercu_departs_retraite.html', context)


# Cette fonction permet de générer le contrat de travail
@login_required(login_url="user_signin")
def generate_contract_pdf(request, id):
    contrat = get_object_or_404(Contrat, id=id)
    par = Parametre_calcul.objects.get(id=1)
    sal_brut = salaire_brut(contrat, par)

    if not contrat:
        return HttpResponse("Le contrat n'existe pas.", status=404)

    type_contrat = contrat.type_contrat.libelle

    if type_contrat not in ['CDD', 'CDI']:
        return HttpResponse("Type de contrat non valide.", status=400)

    contrat_file = contrat.type_contrat.file

    if not contrat_file:
        return HttpResponse("Le fichier du contrat n'existe pas.", status=404)

    if not contrat_file.name.endswith('.docx'):
        return HttpResponse("Le fichier n'est pas au format Word.", status=400)

    doc_content = contrat_file.read()
    doc = Document(io.BytesIO(doc_content))

    Date_Premiere_Periode = contrat.date_debut + timedelta(days=30)

    contrat_dict = {
        "M": "Monsieur",
        "F": "Madame"
    }

    for paragraph in doc.paragraphs:
        if contrat.employe.sexe == 'M':
            paragraph.text = paragraph.text.replace('[APPELATION]', contrat_dict['M'])
            paragraph.text = paragraph.text.replace('[ENFANT]', 'Fils')
        elif contrat.employe.sexe == 'F':
            paragraph.text = paragraph.text.replace('[APPELATION]', contrat_dict['F'])
            paragraph.text = paragraph.text.replace('[ENFANT]', 'Fille')
        paragraph.text = paragraph.text.replace('[NOM_EMPLOYE]', contrat.employe.nom + ' ' + contrat.employe.prenoms)
        paragraph.text = paragraph.text.replace('[NOM_EMPLOYE_PERE]', contrat.employe.nom_pere)
        paragraph.text = paragraph.text.replace('[NOM_EMPLOYE_MERE]', contrat.employe.nom_mere)
        paragraph.text = paragraph.text.replace('[DATE_NAISSANCE]', str(contrat.employe.date_naissance))
        paragraph.text = paragraph.text.replace('[LIEU_NAISSANCE]', contrat.employe.lieu_naissance)
        paragraph.text = paragraph.text.replace('[SITUATION_MATRIMONIALE]', str(contrat.employe.situation_matrimoniale))
        paragraph.text = paragraph.text.replace('[RESIDENCE]', contrat.employe.adresse)
        paragraph.text = paragraph.text.replace('[NUMERO_CNIB]', str(contrat.employe.numero_cnib))
        paragraph.text = paragraph.text.replace('[profession]', contrat.employe.profession)
        paragraph.text = paragraph.text.replace('[TELEPHONE]', str(contrat.employe.telephone))
        paragraph.text = paragraph.text.replace('[PERSONNE_PREVENIR]', contrat.employe.personne_prevenir)
        paragraph.text = paragraph.text.replace('[TELEPHONE_PREVENIR]', str(contrat.employe.telephone_prevenir))
        paragraph.text = paragraph.text.replace('[DATE_DEBUT_CONTRAT]', str(contrat.date_debut))
        paragraph.text = paragraph.text.replace('[SOUS_COUVERT]', contrat.employe.sous_couvert)
        paragraph.text = paragraph.text.replace('[STRUCTURE]', contrat.structure.denomination)
        paragraph.text = paragraph.text.replace('[ADRESSE_STRUCTURE]', contrat.structure.adresse)
        paragraph.text = paragraph.text.replace('[POSTE]', contrat.poste)
        paragraph.text = paragraph.text.replace('[DATE_DEBUT]', str(contrat.date_debut))
        paragraph.text = paragraph.text.replace('[DATE_FIN]', str(contrat.date_fin))
        paragraph.text = paragraph.text.replace('[DATE_PREMIER_PERIODE]', str(Date_Premiere_Periode))
        paragraph.text = paragraph.text.replace('[SALAIRE_BRUT]', str(sal_brut))
        paragraph.text = paragraph.text.replace('[SALAIRE_BASE]', str(contrat.salaire_base))
        paragraph.text = paragraph.text.replace('[INDEMNITE_LOG]', str(contrat.indemnite_logement))
        paragraph.text = paragraph.text.replace('[INDEMNITE_TRANSP]', str(contrat.indemnite_transport))
        paragraph.text = paragraph.text.replace('[INDEMNITE_FONC]', str(contrat.indemnite_fonction))
        paragraph.text = paragraph.text.replace('[SALAIRE_BRUT_LETTRE]', num2words(sal_brut, lang='fr').capitalize())
        paragraph.text = paragraph.text.replace('[SALAIRE_BASE_LETTRE]', num2words(contrat.salaire_base, lang='fr').capitalize())
        paragraph.text = paragraph.text.replace('[INDEMNITE_LOG_L]', num2words(contrat.indemnite_logement, lang='fr').capitalize())
        paragraph.text = paragraph.text.replace('[INDET_TRANP_L]', num2words(contrat.indemnite_transport, lang='fr').capitalize())
        paragraph.text = paragraph.text.replace('[INDETE_FONC_LE]', num2words(contrat.indemnite_fonction, lang='fr').capitalize())

    contrat_rempli_path = f"Contrat_{contrat.employe.nom}_{contrat.employe.prenoms}.docx"
    doc.save(contrat_rempli_path)

    pdf_path = f"Contrat_{contrat.employe.nom}_{contrat.employe.prenoms}.pdf"
    convert_to_pdf(contrat_rempli_path, pdf_path)

    response = HttpResponse(open(pdf_path, 'rb'), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="contrat_{contrat.employe.nom}_{contrat.type_contrat}.pdf"'

    os.remove(contrat_rempli_path)
    os.remove(pdf_path)

    return response

def convert_to_pdf(input_docx, output_pdf):
    c = canvas.Canvas(output_pdf, pagesize=letter)
    width, height = letter
    doc = Document(input_docx)
    y_position = height - 40
    c.setFont("Helvetica", 12)
    for para in doc.paragraphs:
        text = para.text
        lines = text.split('\n')
        for line in lines:
            if len(line) > 90:
                while len(line) > 90:
                    part = line[:90]
                    c.drawString(40, y_position, part)
                    y_position -= 15
                    line = line[90:]
            c.drawString(40, y_position, line)
            y_position -= 15
            if y_position < 40:
                c.showPage()
                y_position = height - 40
                c.setFont("Helvetica", 12)

    c.save()

#Module bnts
@login_required(login_url="user_signin")
def bnts_declaration(request):
    declarations = Contrat.objects.select_related('employe', 'structure').all()
    for declaration in declarations:
        paiement = Paiement.objects.filter(contrat=declaration).order_by('-created_at').first()
        if paiement:
            declaration.base_cnss = paiement.cotitsation_cnss
            declaration.salaire_brut = paiement.salaire_brut
    
    context = {
        'declarations': declarations,
    }
    return render(request, 'calcul/bnts_declaration.html', context)

#Déclaration tpa
@login_required(login_url="user_signin")
def tpa_declaration(request):
    declarations = Contrat.objects.select_related('employe').all()
    for declaration in declarations:
        paiement = Paiement.objects.filter(contrat=declaration).order_by('-created_at').first()
        if paiement:
            declaration.salaire_brut = paiement.salaire_brut
            declaration.tpa = paiement.tpa
    
    context = {
        'declarations': declarations,
    }
    return render(request, 'calcul/tpa_declaration.html', context)

#Déclaration iuts
@login_required(login_url="user_signin")
def iuts_declaration(request):
    declarations = Contrat.objects.select_related('employe').all()
    for declaration in declarations:
        paiement = Paiement.objects.filter(contrat=declaration).order_by('-created_at').first()
        if paiement:
            declaration.salaire_brut = paiement.salaire_brut
            declaration.base_iuts = paiement.salaire_brut  # Ajustez si nécessaire
            declaration.iuts = paiement.iuts
    
    context = {
        'declarations': declarations,
    }
    return render(request, 'calcul/iuts_declaration.html', context)


